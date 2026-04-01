"""
DOCX generator for CSI MasterFormat specification sections.

Applies the style profile defined in CLAUDE_CODE_INSTRUCTIONS.md:

  Element         Font    Size   Bold   Caps   SpaceBefore  SpaceAfter
  ──────────────  ──────  ─────  ─────  ─────  ───────────  ──────────
  Section Title   Arial   14pt   yes    ALL    240 twips    120 twips
  Part Header     Arial   12pt   yes    Mixed  200 twips    120 twips
  Sub-Section     Arial   11pt   yes    ALL    160 twips    60  twips
  Body Text       Arial   10pt   no     —      0   twips    120 twips
  Bullet L1/L2    Arial   10pt   no     —      0   twips    60  twips
  Footer          Arial    9pt   no     —      —            —

  Page: US Letter 8.5"×11", margins top/bottom/right 1", left 1.25"
  Header: project name + bottom border rule
  Footer: "Page N" (left)  |  owner name (right)
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

from .models import BULLET_L1, BULLET_L2, PARAGRAPH, SpecSection

# ── Style constants ───────────────────────────────────────────────────────────

_FONT = "Arial"

# Font sizes
_SZ_SECTION = Pt(14)
_SZ_PART    = Pt(12)
_SZ_SUBSEC  = Pt(11)
_SZ_BODY    = Pt(10)
_SZ_FOOTER  = Pt(9)

# Spacing (twips; 1 pt = 20 twips)
_SP_SEC_BEFORE    = Twips(240)
_SP_SEC_AFTER     = Twips(120)
_SP_PART_BEFORE   = Twips(200)
_SP_PART_AFTER    = Twips(120)
_SP_SUBSEC_BEFORE = Twips(160)
_SP_SUBSEC_AFTER  = Twips(60)
_SP_BODY_BEFORE   = Twips(0)
_SP_BODY_AFTER    = Twips(120)
_SP_BULLET_BEFORE = Twips(0)
_SP_BULLET_AFTER  = Twips(60)

# Page layout
_PAGE_W    = Inches(8.5)
_PAGE_H    = Inches(11.0)
_MARGIN_T  = Inches(1.0)
_MARGIN_B  = Inches(1.0)
_MARGIN_L  = Inches(1.25)
_MARGIN_R  = Inches(1.0)

# Text area width in twips (used for right-tab in footer)
# 8.5 - 1.25 - 1.0 = 6.25 inches → 6.25 * 1440 = 9000 twips
_TEXT_WIDTH_TWIPS = 9000

# Bullet indents
_BULLET1_LEFT    = Inches(0.375)
_BULLET1_HANG    = Inches(-0.25)
_BULLET2_LEFT    = Inches(0.625)
_BULLET2_HANG    = Inches(-0.25)

# Sub-section heading left indent
_SUBSEC_INDENT = Inches(0.0)   # flush-left; visually set off by font size difference


# ── Public API ────────────────────────────────────────────────────────────────

def build_docx(spec: SpecSection, output_path: str | None = None) -> str:
    """
    Build a styled DOCX from *spec* and save it.

    Args:
        spec: Fully-populated SpecSection from the wizard.
        output_path: Override for output file path; falls back to spec.output_path.

    Returns:
        The resolved output file path.
    """
    path = output_path or spec.output_path or "output_spec.docx"

    doc = Document()
    _configure_page(doc)
    _add_header(doc, spec)
    _add_footer(doc, spec)
    _write_body(doc, spec)

    doc.save(path)
    return path


# ── Page setup ────────────────────────────────────────────────────────────────

def _configure_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width    = _PAGE_W
    sec.page_height   = _PAGE_H
    sec.top_margin    = _MARGIN_T
    sec.bottom_margin = _MARGIN_B
    sec.left_margin   = _MARGIN_L
    sec.right_margin  = _MARGIN_R
    sec.different_first_page_header_footer = False


# ── Header ────────────────────────────────────────────────────────────────────

def _add_header(doc: Document, spec: SpecSection) -> None:
    """Project name in header with a bottom-border rule."""
    header = doc.sections[0].header
    header.is_linked_to_previous = False

    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Project name + number on one line
    title_text = spec.project.project_name
    if spec.project.project_number:
        title_text = f"{spec.project.project_name}  —  {spec.project.project_number}"

    run = hp.add_run(title_text)
    run.font.name = _FONT
    run.font.size = Pt(9)
    run.font.bold = True

    _set_paragraph_border(hp, bottom=True)


# ── Footer ────────────────────────────────────────────────────────────────────

def _add_footer(doc: Document, spec: SpecSection) -> None:
    """Page number on the left, owner name on the right."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # "Page " + field
    _add_run(fp, "Page ", _SZ_FOOTER)
    _insert_page_field(fp)

    # Tab → right-aligned owner name
    tab_run = fp.add_run("\t")
    tab_run.font.name = _FONT
    tab_run.font.size = _SZ_FOOTER

    _add_run(fp, spec.project.owner, _SZ_FOOTER)
    _set_right_tab(fp, _TEXT_WIDTH_TWIPS)


# ── Body ──────────────────────────────────────────────────────────────────────

def _write_body(doc: Document, spec: SpecSection) -> None:
    # Remove the empty default paragraph python-docx inserts
    if doc.paragraphs and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # ── Section header block ──────────────────────────────────────────────────
    # Line 1: "SECTION 01 11 00"
    _section_title(doc, f"SECTION {spec.csi_number}", space_after=Twips(0))
    # Line 2: section title (ALL CAPS)
    _section_title(doc, spec.section_title.upper(), space_before=Twips(0))

    # ── Parts ─────────────────────────────────────────────────────────────────
    for part in spec.parts:
        _write_part(doc, part)

    # ── END OF SECTION ────────────────────────────────────────────────────────
    _spacer(doc)
    p = doc.add_paragraph()
    run = p.add_run("END OF SECTION")
    run.font.name = _FONT
    run.font.size = _SZ_BODY
    run.font.bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = _SP_BODY_BEFORE
    p.paragraph_format.space_after  = _SP_BODY_AFTER


def _write_part(doc: Document, part) -> None:
    # "PART 1 – GENERAL"
    label = f"PART {part.part_number} \u2013 {part.title.upper()}"
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.name  = _FONT
    run.font.size  = _SZ_PART
    run.font.bold  = True
    pf = p.paragraph_format
    pf.space_before = _SP_PART_BEFORE
    pf.space_after  = _SP_PART_AFTER
    pf.alignment    = WD_ALIGN_PARAGRAPH.LEFT

    if part.not_used:
        _body(doc, "(NOT USED)")
        return

    for ss in part.sub_sections:
        _write_subsection(doc, ss)


def _write_subsection(doc: Document, ss) -> None:
    # "1.1  SUMMARY"
    label = f"{ss.number}\u00a0\u00a0{ss.title.upper()}"
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.name    = _FONT
    run.font.size    = _SZ_SUBSEC
    run.font.bold    = True
    run.font.all_caps = True
    pf = p.paragraph_format
    pf.space_before  = _SP_SUBSEC_BEFORE
    pf.space_after   = _SP_SUBSEC_AFTER
    pf.left_indent   = _SUBSEC_INDENT
    pf.alignment     = WD_ALIGN_PARAGRAPH.LEFT

    for line in ss.lines:
        if line.kind == PARAGRAPH:
            _body(doc, line.text)
        elif line.kind == BULLET_L1:
            _bullet(doc, line.text, level=1)
        elif line.kind == BULLET_L2:
            _bullet(doc, line.text, level=2)


# ── Paragraph type builders ───────────────────────────────────────────────────

def _section_title(
    doc: Document,
    text: str,
    space_before=_SP_SEC_BEFORE,
    space_after=_SP_SEC_AFTER,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name    = _FONT
    run.font.size    = _SZ_SECTION
    run.font.bold    = True
    run.font.all_caps = True
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after  = space_after
    pf.alignment    = WD_ALIGN_PARAGRAPH.LEFT


def _body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.name = _FONT
        run.font.size = _SZ_BODY
        run.font.bold = bold
    pf = p.paragraph_format
    pf.space_before = _SP_BODY_BEFORE
    pf.space_after  = _SP_BODY_AFTER
    pf.alignment    = WD_ALIGN_PARAGRAPH.LEFT


def _bullet(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()

    # Bullet character + non-breaking space
    bullet_char = "\u2022\u00a0"
    run = p.add_run(bullet_char + text)
    run.font.name = _FONT
    run.font.size = _SZ_BODY

    pf = p.paragraph_format
    pf.space_before = _SP_BULLET_BEFORE
    pf.space_after  = _SP_BULLET_AFTER
    pf.alignment    = WD_ALIGN_PARAGRAPH.LEFT

    if level == 1:
        pf.left_indent        = _BULLET1_LEFT
        pf.first_line_indent  = _BULLET1_HANG
    else:
        pf.left_indent        = _BULLET2_LEFT
        pf.first_line_indent  = _BULLET2_HANG


def _spacer(doc: Document) -> None:
    """Insert a blank body-height paragraph as vertical space."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = _SP_BODY_BEFORE
    p.paragraph_format.space_after  = _SP_BODY_AFTER


# ── XML / run helpers ─────────────────────────────────────────────────────────

def _add_run(paragraph, text: str, size) -> None:
    run = paragraph.add_run(text)
    run.font.name = _FONT
    run.font.size = size


def _set_paragraph_border(p, bottom: bool = False, top: bool = False) -> None:
    """Apply a thin single border to the top or bottom of a paragraph."""
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    for side, flag in (("top", top), ("bottom", bottom)):
        if not flag:
            continue
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"),   "single")
        bdr.set(qn("w:sz"),    "6")       # ¾ pt line
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "000000")
        pBdr.append(bdr)

    pPr.append(pBdr)


def _insert_page_field(paragraph) -> None:
    """Append an auto-updating PAGE field to *paragraph*."""
    for fld_type, instr in (
        ("begin", None),
        (None,    " PAGE "),
        ("end",   None),
    ):
        run = paragraph.add_run()
        run.font.name = _FONT
        run.font.size = _SZ_FOOTER

        if fld_type is not None:
            fldChar = OxmlElement("w:fldChar")
            fldChar.set(qn("w:fldCharType"), fld_type)
            run._r.append(fldChar)
        else:
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = instr
            run._r.append(instrText)


def _set_right_tab(paragraph, position_twips: int) -> None:
    """Add a right-aligned tab stop at *position_twips* from the text margin."""
    pPr  = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position_twips))
    tabs.append(tab)
    pPr.append(tabs)
